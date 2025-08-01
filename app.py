from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from io import BytesIO
import os
import openai
import fitz  # PyMuPDF
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from docx import Document
import re

load_dotenv()
app = Flask(__name__)
CORS(app)

SECTION_LABELS = {
    'objectives': ["objective", "lesson objective", "learning objective"],
    'standards': ["standard", "common core", "missouri standard"],
    'materials': ["materials", "resources needed"],
    'vocabulary': ["vocabulary", "academic vocabulary"],
    'assessment': ["assessment", "formative", "summative"],
    'do_now': ["do now", "bell ringer", "warm-up"],
    'i_do': ["i do", "direct instruction", "teacher modeling"],
    'we_do': ["we do", "guided practice"],
    'you_do_together': ["you do together", "collaborative"],
    'you_do_alone': ["you do alone", "independent practice"],
    'homework': ["homework", "take home tasks"],
    'reflection': ["reflection", "teacher notes"],
}

def generate_lesson_content(prompt):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a lesson plan generator. Use the format {{section}}: content."},
            {"role": "user", "content": f"Create a full lesson plan for: {prompt}. Include objectives, standards, vocabulary, materials, I do, we do, you do together, you do alone, assessment, reflection."}
        ],
        temperature=0.3,
        max_tokens=1800
    )
    return response.choices[0].message.content

def parse_lesson_content(content):
    pattern = r'\{\{(.*?)\}\}:(.*?)(?=\n\{\{|\Z)'
    return {k.strip().lower(): v.strip() for k, v in re.findall(pattern, content, re.DOTALL)}

def match_docx_sections(doc_stream, sections):
    doc = Document(doc_stream)
    for para in doc.paragraphs:
        p_text = para.text.lower().strip()
        for key, labels in SECTION_LABELS.items():
            if key in sections and any(lbl in p_text for lbl in labels):
                idx = doc.paragraphs.index(para)
                doc.paragraphs.insert(idx + 1, doc.add_paragraph(sections[key]))
                break
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def detect_pdf_labels(pdf_stream, sections):
    pdf_stream.seek(0)
    doc = fitz.open(stream=pdf_stream.read(), filetype="pdf")
    positions = {}
    for page_number in range(len(doc)):
        blocks = doc[page_number].get_text("blocks")
        for block in blocks:
            text = block[4].lower().strip()
            for key, labels in SECTION_LABELS.items():
                if key in sections and any(lbl in text for lbl in labels):
                    if key not in positions:
                        x, y = block[0] + 100, block[1]
                        positions[key] = (x, y, page_number)
    return positions

def fill_pdf(pdf_stream, sections, positions):
    pdf_stream.seek(0)
    reader = PyPDF2.PdfReader(pdf_stream)
    writer = PyPDF2.PdfWriter()
    width = float(reader.pages[0].mediabox.width)
    height = float(reader.pages[0].mediabox.height)
    packets = [BytesIO() for _ in reader.pages]
    canvases = []

    for i in range(len(reader.pages)):
        canvases.append(canvas.Canvas(packets[i], pagesize=(width, height)))

    for key, (x, y, page_idx) in positions.items():
        if key in sections:
            can = canvases[page_idx]
            y = height - y
            lines = simpleSplit(sections[key], "Helvetica", 9, 400)
            can.setFont("Helvetica", 9)
            for line in lines:
                can.drawString(x, y, line)
                y -= 12

    filled = BytesIO()
    for i, base_page in enumerate(reader.pages):
        canvases[i].save()
        packets[i].seek(0)
        overlay = PyPDF2.PdfReader(packets[i])
        base_page.merge_page(overlay.pages[0])
        writer.add_page(base_page)

    writer.write(filled)
    filled.seek(0)
    return filled

def fallback_pdf_write(sections, reader):
    writer = PyPDF2.PdfWriter()
    width = float(reader.pages[0].mediabox.width)
    height = float(reader.pages[0].mediabox.height)
    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=(width, height))
    y = height - 50
    can.setFont("Helvetica", 9)

    for key, value in sections.items():
        lines = simpleSplit(f"{key.upper()}: {value}", "Helvetica", 9, 450)
        for line in lines:
            can.drawString(50, y, line)
            y -= 12
        y -= 8
        if y < 50:
            break

    can.save()
    packet.seek(0)
    overlay = PyPDF2.PdfReader(packet)
    reader.pages[0].merge_page(overlay.pages[0])
    writer.add_page(reader.pages[0])
    for pg in reader.pages[1:]:
        writer.add_page(pg)

    filled = BytesIO()
    writer.write(filled)
    filled.seek(0)
    return filled

@app.route('/process', methods=['POST'])
def process():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files['file']
    prompt = request.form.get('prompt', '')
    ext = os.path.splitext(file.filename)[1].lower()

    try:
        content = generate_lesson_content(prompt)
        sections = parse_lesson_content(content)

        if ext == ".docx":
            filled = match_docx_sections(file.stream, sections)
            return send_file(filled, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="filled_lesson_plan.docx")

        elif ext == ".pdf":
            file.stream.seek(0)
            positions = detect_pdf_labels(file.stream, sections)
            file.stream.seek(0)
            if positions:
                filled = fill_pdf(file.stream, sections, positions)
            else:
                reader = PyPDF2.PdfReader(file.stream)
                filled = fallback_pdf_write(sections, reader)
            return send_file(filled, mimetype="application/pdf", as_attachment=True, download_name="filled_lesson_plan.pdf")

        else:
            return jsonify({"error": "Unsupported file type."}), 400
    except Exception as e:
        print("Error:", str(e))
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host="0.0.0.0", port=5000, debug=True)
